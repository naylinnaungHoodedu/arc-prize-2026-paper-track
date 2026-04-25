from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "ARC_Prize_2026_Writeup_Final.docx"

LIGHT_BLUE = "EAF2F8"
GRAY = "F4F6F7"


def shade_paragraph(paragraph, fill: str = GRAY, border: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border:
        borders = p_pr.find(qn("w:pBdr"))
        if borders is None:
            borders = OxmlElement("w:pBdr")
            p_pr.append(borders)
        for edge in ("top", "left", "bottom", "right"):
            element = borders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
            element.set(qn("w:space"), "4")
            element.set(qn("w:color"), border)


def add_run(paragraph, text: str, bold: bool = False, italic: bool = False) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic


def add_callout(doc: Document, title: str, body: str, fill: str = GRAY) -> None:
    p = doc.add_paragraph()
    shade_paragraph(p, fill=fill, border="B7C9E2")
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    add_run(p, title + ": ", bold=True)
    add_run(p, body)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(4)


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.05


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run(text)


def add_labeled_bullet(doc: Document, label: str, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    add_run(paragraph, label + ": ", bold=True)
    add_run(paragraph, text)


def add_number(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run(text)


def add_code_box(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    shade_paragraph(p, fill="F8FAFC", border="C8D6EA")
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.16)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(31, 31, 31)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(8.9)
    normal.font.color.rgb = RGBColor(31, 31, 31)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.05

    for style_name, size in (("Heading 1", 12.8), ("Heading 2", 10.8), ("Heading 3", 10.0)):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 56, 100)
        style.font.size = Pt(size)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(8.8)
        style.paragraph_format.space_after = Pt(2)

    doc.core_properties.title = "ARC Prize 2026 Paper Track Writeup"
    doc.core_properties.subject = "Hybrid Theory-Based Agents for ARC-AGI-3"
    doc.core_properties.author = "Nay Linn Aung"
    doc.core_properties.keywords = "ARC-AGI-3, HTBA, MDL, Bayesian reasoning, Kaggle Paper Track"


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("Hybrid Theory-Based Agents for ARC-AGI-3")
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 56, 100)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("Bayesian skill-acquisition under Core-Knowledge priors")
    run.italic = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = RGBColor(63, 63, 63)

    rows = [
        ("Competition", "ARC Prize 2026 - Paper Track"),
        ("Author", "Nay Linn Aung, Hood College M.S. Computer Science"),
        ("Kaggle team", "Nay Linn Aung"),
        ("Source of truth", "ARC_Prize_2026_Writeup_Final.md"),
    ]
    for left, right in rows:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        add_run(p, left + ": ", bold=True)
        add_run(p, right)

    add_callout(
        doc,
        "Submission boundary",
        "This document is the final writeup derived from the local final Markdown body. "
        "It describes the proof-of-concept package, not a leaderboard-optimized agent, "
        "and it does not claim official ARC-AGI-3 scores that were not produced locally.",
        fill="FFF2CC",
    )
    add_callout(
        doc,
        "Rubric posture",
        "The writeup optimizes conceptual clarity, reasoning logic, documentation quality, "
        "offline reproducibility, and auditable agent behavior.",
        fill=LIGHT_BLUE,
    )


def add_assumption_table(doc: Document) -> None:
    add_callout(
        doc,
        "Assumption ledger",
        "The final writeup explicitly cites interpretation boundaries instead of hiding them in prose.",
        fill=LIGHT_BLUE,
    )
    rows = [
        (
            "All files means every readable file in the current working directory.",
            "The blueprint, draft writeup, notebook, package code, tests, and audit artifacts are treated as local evidence.",
        ),
        (
            "The blueprint defines the intended architecture, narrative, and evaluation criteria.",
            "HTBA, Core-Knowledge priors, MDL, EIG, offline execution, and auditability are non-negotiable design anchors.",
        ),
        (
            "ARC-AGI-3 requires generalization, not memorization or brute-force search.",
            "The DSL is intentionally small and complexity-penalized; no public-task-specific tricks are introduced.",
        ),
        (
            "The Kaggle notebook is a proof of concept, not a leaderboard-optimized solution.",
            "The implementation verifies reasoning primitives and traceability; performance claims remain bounded.",
        ),
        (
            "Official ARC-AGI-3 resources are absent locally.",
            "The notebook fails fast until the official toolkit is supplied by Kaggle, a local install, or an attached input.",
        ),
    ]
    for idx, (left, right) in enumerate(rows, start=1):
        add_labeled_bullet(doc, f"A{idx}", left + " Applied as follows: " + right)


def add_architecture_table(doc: Document) -> None:
    rows = [
        ("C1 Frame Encoder", "Connected components with color, size, bbox, centroid, shape signature, and motion delta.", "Uses inspectable objectness. The proof-of-concept does not invent missing CNN weights."),
        ("C2 Action Probe", "Reads alive actions from official metadata and normalizes RESET, ACTION1-7, including coordinate clicks.", "Avoids wasting actions on dead controls and keeps action semantics toolkit-independent."),
        ("C3 Hypothesis Beam", "Maintains K=64 symbolic rule programs with MDL prior and posterior normalization.", "Makes rule choice auditable and controls search growth."),
        ("C4 Goal Inferer", "Stores only frame deltas observed with WIN.", "Prevents reward claims before evidence exists."),
        ("C5 Planner", "Uses EIG in explore mode and bounded depth-6 expected progress in exploit mode.", "Separates information gathering from goal-directed execution."),
        ("C6 Memory", "SHA-256 content-addressed Core-Knowledge signatures.", "Supports reproducible reuse without hidden state."),
        ("C7 Harness", "Unit tests, notebook checks, static audit, scorecard, and trace outputs.", "Connects claims to evidence."),
    ]
    for component, role, rationale in rows:
        add_labeled_bullet(doc, component, role + " Rationale: " + rationale)


def add_risk_table(doc: Document) -> None:
    rows = [
        ("Overfitting reasoning to ARC patterns", "Keep primitives Core-Knowledge-aligned, apply MDL penalties, and require task-family validation when official resources are available."),
        ("Reasoning appears post-hoc", "Trace records store observation, hypothesis, transformation, and pending validation before the next frame is observed."),
        ("Notebook too complex for reviewers", "Use progressive disclosure: narrative first, public API second, execution cells last."),
        ("Unverifiable capability claims", "Document deterministic perception, minimal DSL coverage, and absence of local official scores."),
    ]
    for left, right in rows:
        add_labeled_bullet(doc, left, right)


def add_evidence_table(doc: Document) -> None:
    rows = [
        ("Unit reasoning checks", "20 passed, 2 integration tests deselected without official resources.", "tests/"),
        ("Notebook schema and code cells", "Valid notebook with 14 cells and 4 executable code cells.", "HTBA_ARC_AGI3_PaperTrack.ipynb"),
        ("Static offline audit", "PASS: seed, notebook, package, cover present; no findings.", "out/audit.html"),
        ("Official ARC run", "Not run locally because the official toolkit is absent.", "Preflight gate in htba/arc_adapter.py"),
    ]
    for item, status, artifact in rows:
        add_labeled_bullet(doc, item, status + f" Artifact: {artifact}.")


def build_doc() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("ARC Prize 2026 Paper Track | HTBA | MIT-0 code / CC0-1.0 docs")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(90, 90, 90)

    add_heading(doc, "Abstract")
    add_body(
        doc,
        "ARC-AGI-3 changes ARC from static grid induction into interactive skill acquisition. "
        "The agent chooses what to observe, acts under an action budget, and is judged by Relative "
        "Human Action Efficiency: RHAE = (human_actions / AI_actions)^2, capped at 1.0. "
        "The Hybrid Theory-Based Agent (HTBA) follows the local blueprint by combining object-centric "
        "perception, Bayesian/MDL symbolic hypotheses, WIN-linked goal inference, information-gain "
        "planning, and deterministic audit artifacts. Its claim is intentionally bounded: the submitted "
        "package verifies a reasoning architecture and trace discipline for the Paper Track; it does not "
        "claim unmeasured leaderboard performance or hidden learned capabilities."
    )

    add_heading(doc, "Assumptions and Scope")
    add_assumption_table(doc)

    add_heading(doc, "1. Problem Framing")
    add_body(
        doc,
        "ARC-AGI-3 is not a larger version of ARC-AGI-1/2. A static solver receives examples and returns "
        "an output grid. An ARC-AGI-3 agent receives a frame, selects an action, observes the next frame, "
        "and must infer both dynamics and goal structure. The scoring function makes inefficient exploration "
        "visible: repeated probing is penalized quadratically, so brute force is a poor conceptual fit even "
        "when it sometimes solves a level."
    )
    add_body(
        doc,
        "The writeup therefore prioritizes clarity of reasoning over raw score. The intended reviewer can "
        "inspect what the agent saw, what rule family it considered, what transformation it predicted, and "
        "how the observation validated or weakened that hypothesis."
    )

    add_heading(doc, "2. Reasoning Framework")
    add_body(
        doc,
        "HTBA maintains a posterior over symbolic transition programs. Each program is drawn from a compact "
        "DSL grounded in Core-Knowledge priors: objectness, geometry, numerosity, action-conditioned causality, "
        "and simple object transformations. The prior penalizes description length, so shorter explanations "
        "win when evidence is equal."
    )
    add_code_box(
        doc,
        [
            "Posterior: P(h | history) proportional to P(history | h) * exp(-L(h) / tau)",
            "Likelihood: product over transitions P(next_frame | frame, action, h)",
            "Explore:   choose argmax_a EIG(a)",
            "EIG(a):    expected posterior entropy reduction after trying action a",
            "Exploit:   choose bounded expected-progress action after entropy is low",
        ],
    )
    add_body(
        doc,
        "This separation matters for auditability. Observation is the encoded object set. Hypothesis is the "
        "current posterior beam. Transformation is the MAP prediction and selected action. Validation is written "
        "only after the next frame is observed. The trace schema prevents the explanation from being silently "
        "rewritten after the result is known."
    )

    add_heading(doc, "3. Architecture and Implementation Evidence")
    add_architecture_table(doc)
    add_body(
        doc,
        "One deviation from the broad blueprint is explicit. The blueprint reserves C1 for object-centric "
        "perception and discusses a small CNN as a possible front end. The local final package instantiates C1 "
        "as deterministic connected-component extraction because no trained CNN weights are present in the "
        "files. This is documented as an implementation boundary, not masked as a learned capability."
    )

    add_heading(doc, "4. Algorithmic Flow")
    add_code_box(
        doc,
        [
            "preflight official ARC toolkit resources",
            "create Arcade(operation_mode=OperationMode.COMPETITION)",
            "iterate available official ARC-AGI-3 environments",
            "for each environment:",
            "    reset agent and environment",
            "    encode frame as objects",
            "    initialize or update the MDL posterior beam",
            "    if posterior entropy is high: act by EIG",
            "    else: act by bounded expected progress under WIN-linked deltas",
            "    observe next frame and update validation",
            "write out/scorecard.json, out/reasoning_trace.json, and out/audit.html",
        ],
    )
    add_body(
        doc,
        "The public API remains intentionally small: reset(game), act(frame), observe(action, next_frame, win), "
        "scorecard(), and reasoning_trace(). The notebook supplies the execution boundary; the Python package "
        "supplies the reusable reasoning machinery."
    )

    add_heading(doc, "5. Reproducibility and Evaluation")
    add_body(
        doc,
        "The package is offline-first with respect to hosted AI and runtime installation. It uses a fixed seed, "
        "0xA6C16E26, rejects runtime installation cells, and requires the official ARC-AGI-3 toolkit through "
        "the Kaggle runtime, ARC_AGI_TOOLKIT_DIR, or recognized Kaggle inputs. The optional socket guard is for "
        "local audits only and is not enabled during official competition-mode execution because ARC-managed "
        "toolkit endpoints must remain reachable. There is no synthetic fallback on the official execution path."
    )
    add_evidence_table(doc)
    add_body(
        doc,
        "Accuracy in this Paper Track context is the match between claims and runnable evidence. The current "
        "local evidence validates reasoning primitives, determinism, notebook structure, and offline auditability. "
        "Official task-level RHAE should be reported only after the official toolkit is attached and "
        "the integration-marked run succeeds."
    )

    add_heading(doc, "6. Limitations and Risk Controls")
    add_risk_table(doc)
    add_body(
        doc,
        "The remaining limitation is coverage. A minimal DSL cannot express every possible ARC-AGI-3 environment. "
        "The design choice is deliberate: add primitives only when they have a Core-Knowledge rationale, an explicit "
        "description length, and unit-level validation. That discipline is more valuable for the Paper Track than "
        "a larger opaque rule bank."
    )

    add_heading(doc, "7. Why This Generalizes")
    add_body(
        doc,
        "The generalization claim is architectural, not empirical overreach. The inference loop is reusable whenever "
        "state can be encoded as finite objects, actions change those objects, and sparse success signals can be "
        "linked to frame deltas. The DSL may change by domain; the audit discipline, MDL pressure, and separation of "
        "observation, hypothesis, transformation, and validation do not."
    )

    add_heading(doc, "References")
    references = [
        "ARC Prize Foundation. ARC-AGI-3: A New Challenge for Frontier Agentic Intelligence. arXiv:2603.24621, 2026.",
        "Chollet, F. On the Measure of Intelligence. arXiv:1911.01547, 2019.",
        "Tsividis, P. et al. Human-Level Reinforcement Learning through Theory-Based Modeling, Exploration, and Planning. arXiv:2107.12544, 2021.",
        "Grunwald, P. The Minimum Description Length Principle. MIT Press, 2007.",
        "ARC_Prize_2026_Writeup_Final.md. Final Kaggle Paper Track writeup body, April 25, 2026.",
    ]
    for ref in references:
        add_number(doc, ref)

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build_doc()
